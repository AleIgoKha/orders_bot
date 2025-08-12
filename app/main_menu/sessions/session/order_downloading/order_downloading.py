import re
from aiogram import F, Router, Bot
from aiogram.types import BufferedInputFile, CallbackQuery
from aiogram.fsm.context import FSMContext
from aiogram.enums import ChatAction
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from io import BytesIO
from unidecode import unidecode

import app.main_menu.sessions.session.order_downloading.keyboard as kb
from app.com_func import group_orders_items, vacc_price_counter
from app.database.requests import get_orders_items, get_session, get_orders, get_session_items_stats

order_downloading = Router()


@order_downloading.callback_query(F.data == 'session_downloads')
async def session_downloads_handler(callback: CallbackQuery):
    await callback.message.edit_text(text='❓ <b>ВЫБЕРИТЕ ПУНКТ ДЛЯ ЗАГРУЗКИ</b> ❓',
                                     reply_markup=kb.session_downloads_menu,
                                     parse_mode='HTML')
    

# загружаем список заказов
@order_downloading.callback_query(F.data == 'download_orders')
async def download_orders_handlers(callback: CallbackQuery, state: FSMContext, bot: Bot):
    data = await state.get_data()
    
    # Загружаем все заказы сессии
    session_id = data['session_id']
    session_data = await get_session(session_id=session_id)
    session_name = session_data.session_name
    orders_items = await get_orders_items(session_id=session_id)
    orders_items_data = group_orders_items(orders_items)

    order_items_data = orders_items_data[::-1]
    
    doc = Document()
    
    # устанавливаем разметку страницы
    section = doc.sections[0]
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1)
    
    # устанавливаем стили для документа
    style = doc.styles['Heading 1']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(16)
    font.color.rgb = RGBColor(0, 0, 0)

    style = doc.styles['Heading 3']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0, 0, 0)

    # наполняем документ
    doc.add_heading(f"СПИСОК ЗАКАЗОВ СЕССИИ - {session_name}", level=1)
    for order_items_data in orders_items_data:
        order_number = order_items_data['order_number']
        client_name = order_items_data['client_name']
        
        doc.add_heading(f'{client_name} №{order_number}', level=3)
        p = doc.add_paragraph()
        
        items_list = [item for item in order_items_data.keys() if item.startswith('item_')]
        total_price = 0
        
        if items_list: # Проверяем пуст ли заказ
            for item in items_list:
                item_name = order_items_data[item]['item_name']
                item_price = float(order_items_data[item]['item_price'])
                item_qty = float(order_items_data[item]['item_qty'])
                item_unit = order_items_data[item]['item_unit']
                item_qty_fact = float(order_items_data[item]['item_qty_fact'])
                item_vacc = order_items_data[item]['item_vacc']
                
                if item_vacc:
                    item_vacc = ' (вак. уп.)'
                else:
                    item_vacc = ''
                    
                p.add_run(f'{item_name}{item_vacc} / ').bold = True
                
                if item_unit == 'кг': # Переводим килограмы в граммы
                    p.add_run(f'Заказано - {int(item_qty * 1000)} {item_unit[-1]} / ')
                    p.add_run(f'Взвешено - {int(item_qty_fact * 1000)} {item_unit[-1]} / ')
                else:
                    p.add_run(f'Заказано - {int(item_qty)} {item_unit} / ')
                    p.add_run(f'Взвешено - {int(item_qty_fact)} {item_unit} / ')
                # Рассчитываем стоимость всключая вакуум
                
                # считаем стоимость вакуума
                vacc_price = vacc_price_counter(item_vacc,
                                                item_qty_fact,
                                                item_unit)

                item_price = round(item_qty_fact * float(item_price) + vacc_price)
                total_price += item_price
                
                run = p.add_run(f'Стоимость - {item_price} р')
                run.add_break()
                    
        else:
            run = p.add_run('Заказ пуст ')
            run.add_break()
        

        
        order_disc = order_items_data['order_disc']
        if order_disc > 0:
            p.add_run('Размер скидки - ').bold = True
            run = p.add_run(f'{order_disc} %')
            run.add_break()
        
        p.add_run(f'К ОПЛАТЕ - {round(total_price * ((100 - order_disc) / 100))} р').bold =  True
        run = p.add_run()
        
        order_note = order_items_data['order_note']
        if order_note:
            run.add_break()
            p.add_run('Комментарий к заказу:').bold = True
            run = p.add_run()
            run.add_break()
            p.add_run(f'{order_note}')
        
        p.paragraph_format.line_spacing = Pt(14)
    
    file_bytes = BytesIO()
    doc.save(file_bytes)
    file_bytes.seek(0)
    
    filename = f'{unidecode(session_name)}_orders'.replace(' ', '_').replace('-', '_').lower()
    filename = re.sub(r"[^a-z0-9_]", "", filename)
    
    await callback.bot.send_chat_action(chat_id=data['chat_id'], action=ChatAction.UPLOAD_DOCUMENT)
    
    # удаляем все id всех сообщений в чате
    redis = bot.redis
    chat_id = callback.message.chat.id
    key = f"chat:{chat_id}:messages"
    message_ids = await redis.lrange(key, 0, -1)  # Get all stored message IDs

    for msg_id in message_ids:
        try:
            await bot.delete_message(chat_id, int(msg_id))
        except Exception:
            pass  # Ignore errors (e.g., message already deleted)

    await redis.delete(key)  # Clear the stored list
    await redis.close()

    # печатаем сообщение и сохраняем его id
    sent_message = await callback.bot.send_document(chat_id=data['chat_id'],
                                     document=BufferedInputFile(file=file_bytes.read(), filename=f'{filename}.docx'),
                                     reply_markup=kb.back_from_order_download)
    message_id = sent_message.message_id
    await redis.rpush(f"chat:{chat_id}:messages", message_id)
    await redis.close()



# Скачать статистику в виде docx файла
@order_downloading.callback_query(F.data == 'stats_download')
async def stats_download_handler(callback: CallbackQuery, state: FSMContext, bot: Bot):
    # Запрашиваем данные для всех заказов сессии и данные сессии
    data = await state.get_data()
    session_id = data['session_id']
    session_data = await get_session(session_id=session_id)
    session_name = session_data.session_name
    orders_data = await get_orders(session_id)
    items_stats = await get_session_items_stats(session_id)
    
    doc = Document()
    
    # устанавливаем разметку страницы
    section = doc.sections[0]
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1)
    
    # устанавливаем стили для документа
    style = doc.styles['Heading 1']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(16)
    font.color.rgb = RGBColor(0, 0, 0)

    style = doc.styles['Heading 3']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0, 0, 0)

    # наполняем документ
    doc.add_heading(f"СТАТИСТИКА СЕССИИ - {session_name}", level=1)
    
    p = doc.add_paragraph()
    
    p.add_run(f'Общее количество заказов - ')
    p.add_run(f'{len(orders_data)}').bold = True
    
    # межстрочные интервалы
    p.paragraph_format.line_spacing = Pt(14)
    
    est_revenue = 0
    exp_revenue = 0
    for item_stats in items_stats:
        p = doc.add_paragraph()
        p.add_run(f'{item_stats[0]}').bold = True
        run = p.add_run()
        run.add_break()
        
        if item_stats[1] == 'кг':
            run = p.add_run(f'Всего заказано - {round(item_stats[2], 3)} {item_stats[1]}')
            run.add_break()
            run = p.add_run(f'Всего взвешено - {round(item_stats[3], 3)} {item_stats[1]}')
            run.add_break() 
        else:
            run = p.add_run(f'Всего заказано - {round(item_stats[2])} {item_stats[1]}')
            run = p.add_run(f'Всего взвешено - {round(item_stats[3])} {item_stats[1]}')
            run.add_break() 
            
        run = p.add_run(f'Заказано в вакууме - {round(item_stats[6])} шт.')
        run.add_break()
                
        est_revenue += item_stats[4] # по заказанным количествам
        exp_revenue += item_stats[5] # по фактическим количествам
        
        # межстрочные интервалы
        p.paragraph_format.line_spacing = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run(f'Ожидаемая выручка (без учета вак. уп.)').bold = True
    run = p.add_run()
    run.add_break()
    run = p.add_run(f'По зазанному количеству - {round(est_revenue)} р')
    run.add_break()
    p.add_run(f'По взвешенному количеству - {round(exp_revenue)} р')
    
    # межстрочные интервалы
    p.paragraph_format.line_spacing = Pt(14)
    
    file_bytes = BytesIO()
    doc.save(file_bytes)
    file_bytes.seek(0)
    
    
    # Убираем слэши, пробелы и делаем символы маленькими, переводим в нижний регистр и убираем специальные символы
    filename = f'{unidecode(session_name)}_stats'.replace(' ', '_').replace('-', '_').lower()
    filename = re.sub(r"[^a-z0-9_]", "", filename)
    
    await callback.bot.send_chat_action(chat_id=data['chat_id'], action=ChatAction.UPLOAD_DOCUMENT)
    
    # удаляем все id всех сообщений в чате
    redis = bot.redis
    chat_id = callback.message.chat.id
    key = f"chat:{chat_id}:messages"
    message_ids = await redis.lrange(key, 0, -1)  # Get all stored message IDs

    for msg_id in message_ids:
        try:
            await bot.delete_message(chat_id, int(msg_id))
        except Exception:
            pass  # Ignore errors (e.g., message already deleted)

    await redis.delete(key)  # Clear the stored list
    await redis.close()
    
    # печатаем сообщение и сохраняем его id
    sent_message = await callback.bot.send_document(chat_id=data['chat_id'],
                                     document=BufferedInputFile(file=file_bytes.read(), filename=f'{filename}.docx'),
                                     reply_markup=kb.back_from_order_download)
    message_id = sent_message.message_id
    await redis.rpush(f"chat:{chat_id}:messages", message_id)
    await redis.close()