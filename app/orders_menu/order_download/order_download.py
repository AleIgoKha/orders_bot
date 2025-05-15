from aiogram import F, Router
from aiogram.types import BufferedInputFile, CallbackQuery
from aiogram.fsm.context import FSMContext
from aiogram.enums import ChatAction
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from io import BytesIO

import app.orders_menu.order_download.keyboard as kb
from app.com_func import group_orders_items
from app.database.requests import get_orders_items, get_session

order_download = Router()


@order_download.callback_query(F.data == 'download_orders')
async def download_orders_handlers(callback: CallbackQuery, state: FSMContext):
    data = await state.get_data()
    
    # Загружаем все заказы сессии
    session_id = data['session_id']
    session_data = await get_session(session_id=session_id)
    session_date = str(session_data.session_date).split(' ')[0]
    session_place = session_data.session_place
    orders_items = await get_orders_items(session_id=session_id)
    orders_items_data = group_orders_items(orders_items)

    order_items_data = orders_items_data[::-1]
    
    doc = Document()
    
    # устанавливаем разметку страницы
    section = doc.sections[0]
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1)
    
    # устанавливаем стили для документа
    style = doc.styles['Heading 1']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(16)
    font.color.rgb = RGBColor(0, 0, 0)

    style = doc.styles['Heading 3']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0, 0, 0)

    # наполняем документ
    doc.add_heading(f"{session_date} - {session_place}", level=1)
    for order_items_data in orders_items_data:
        order_number = order_items_data['order_number']
        client_name = order_items_data['client_name']
        
        doc.add_heading(f'{client_name} №{order_number}', level=3)
        p = doc.add_paragraph()
        
        items_list = [item for item in order_items_data.keys() if item.startswith('item_')]
        total_price = 0
        
        if items_list: # Проверяем пуст ли заказ
            for item in items_list:
                item_name = order_items_data[item]['item_name']
                item_price = order_items_data[item]['item_price']
                item_qty = order_items_data[item]['item_qty']
                item_unit = order_items_data[item]['item_unit']
                item_qty_fact = order_items_data[item]['item_qty_fact']
                item_vacc = order_items_data[item]['item_vacc']
                
                if item_vacc:
                    item_vacc = ' (вак. уп.)'
                else:
                    item_vacc = ''
                    
                p.add_run(f'{item_name}{item_vacc} / ').bold = True
                
                if item_unit == 'кг': # Переводим килограмы в граммы
                    p.add_run(f'Заказано - {int(item_qty * 1000)} {item_unit[-1]} / ')
                    p.add_run(f'Взвешено - {int(item_qty_fact * 1000)} {item_unit[-1]} / ')
                else:
                    item_qty_fact = item_qty
                    p.add_run(f'Заказано - {int(item_qty)} {item_unit} / ')
                # Рассчитываем стоимость всключая вакуум
                
                if item_vacc:
                    if item_qty_fact < 200:
                        vacc_price = 5
                    elif 200 <= item_qty_fact < 300:
                        vacc_price = 6
                    elif 300 <= item_qty_fact:
                        vacc_price = (item_qty_fact * 2) // 100
                else:
                    vacc_price = 0

                item_price = round(item_qty_fact * float(item_price)) + vacc_price
                total_price += item_price
                
                run = p.add_run(f'Стоимость - {item_price} р')
                run.add_break()
                    
        else:
            run = p.add_run('Заказ пуст ')
            run.add_break()
        

        
        order_disc = order_items_data['order_disc']
        if order_disc > 0:
            p.add_run('Размер скидки - ').bold = True
            run = p.add_run(f'{order_disc} %')
            run.add_break()
        
        p.add_run(f'К ОПЛАТЕ - {round(total_price * ((100 - order_disc) / 100))} р').bold =  True
        run = p.add_run()
        
        order_note = order_items_data['order_note']
        if order_note:
            run.add_break()
            p.add_run('Комментарий к заказу:').bold = True
            run = p.add_run()
            run.add_break()
            p.add_run(f'{order_note}')
        
        p.paragraph_format.line_spacing = Pt(14)
    
    file_bytes = BytesIO()
    doc.save(file_bytes)
    file_bytes.seek(0)
    
    
    await callback.message.delete()
    await callback.bot.send_chat_action(chat_id=data['chat_id'], action=ChatAction.UPLOAD_DOCUMENT)
    await callback.bot.send_document(chat_id=data['chat_id'],
                                     document=BufferedInputFile(file=file_bytes.read(), filename='generated.docx'),
                                     reply_markup=kb.back_from_order_download)